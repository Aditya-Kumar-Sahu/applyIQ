from __future__ import annotations

from io import BytesIO

import fitz
from docx import Document
import structlog

from app.core.logging_safety import bytes_snapshot, log_debug, log_exception


logger = structlog.get_logger(__name__)


class FileExtractionService:
    ALLOWED_EXTENSIONS = {".pdf", ".docx"}

    def extract_text(self, filename: str, content: bytes) -> str:
        lower_name = filename.lower()
        log_debug(
            logger,
            "resume_extract.extract_text.start",
            filename=filename,
            content=bytes_snapshot(content),
        )
        try:
            if lower_name.endswith(".pdf"):
                log_debug(logger, "resume_extract.extract_text.route", extractor="pdf")
                return self._extract_pdf_text(content)

            if lower_name.endswith(".docx"):
                log_debug(logger, "resume_extract.extract_text.route", extractor="docx")
                return self._extract_docx_text(content)

            raise ValueError("Unsupported resume format")
        except Exception as error:
            log_exception(
                logger,
                "resume_extract.extract_text.failed",
                error,
                filename=filename,
                content=bytes_snapshot(content),
            )
            raise

    def _extract_pdf_text(self, content: bytes) -> str:
        document = fitz.open(stream=content, filetype="pdf")
        page_texts: list[str] = []
        for page in document:
            blocks = page.get_text("blocks")
            if blocks:
                ordered_blocks = sorted(blocks, key=lambda block: (block[1], block[0]))
                block_text = "\n".join(block[4].strip() for block in ordered_blocks if block[4].strip())
                page_texts.append(block_text)
                continue
            page_texts.append(page.get_text("text"))

        text = "\n".join(page.strip() for page in page_texts if page.strip()).strip()
        log_debug(logger, "resume_extract.extract_pdf.complete", pages=len(page_texts), output_length=len(text))
        return text

    def _extract_docx_text(self, content: bytes) -> str:
        document = Document(BytesIO(content))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        table_rows: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_rows.append(" | ".join(cells))

        chunks = [*paragraphs, *table_rows]
        text = "\n".join(chunk for chunk in chunks if chunk).strip()
        log_debug(
            logger,
            "resume_extract.extract_docx.complete",
            paragraphs=len(paragraphs),
            table_rows=len(table_rows),
            output_length=len(text),
        )
        return text
