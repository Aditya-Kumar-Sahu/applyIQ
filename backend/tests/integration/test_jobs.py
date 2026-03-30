from __future__ import annotations

from io import BytesIO
from pathlib import Path

import anyio
from docx import Document
from fastapi.testclient import TestClient
import pytest
from sqlalchemy.exc import IntegrityError
from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine

from app.core.config import Settings
from app.main import create_app
from app.models.base import Base
from app.models.job import Job
from app.services.embedding_service import EmbeddingService


def test_jobs_endpoints_rank_filter_and_search_results(tmp_path: Path) -> None:
    settings = Settings(
        environment="test",
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'jobs.db'}",
        redis_url="redis://localhost:6394/0",
        jwt_secret_key="test-jwt-secret-key-with-32-characters",
        fernet_secret_key="wWKJg6WVKwwhFVWG2yt30YIOCwVDDDeWGPAHDLcGRID=",
        encryption_pepper="pepper-for-tests",
    )

    async def healthy_reporter() -> dict[str, str]:
        return {"status": "ok", "db": "up", "redis": "up"}

    app = create_app(settings=settings, health_reporter=healthy_reporter)

    with TestClient(app) as client:
        anyio.run(_create_all_tables, app.state.database.engine)
        _register_and_prepare_resume(client)
        anyio.run(_seed_jobs, app.state.database.engine)

        jobs_response = client.get("/api/v1/jobs")

        assert jobs_response.status_code == 200
        jobs_payload = jobs_response.json()["data"]
        assert jobs_payload["total"] == 2
        assert jobs_payload["items"][0]["title"] == "Senior ML Engineer"
        assert jobs_payload["items"][0]["recommendation"] in {"strong_match", "good_match"}
        assert jobs_payload["items"][0]["match_score"] > jobs_payload["items"][1]["match_score"]
        assert jobs_payload["items"][0]["score_breakdown"]["location_match"] == 1.0
        assert jobs_payload["items"][0]["score_breakdown"]["salary_alignment"] == 1.0
        assert "Python" in jobs_payload["items"][0]["matched_skills"]

        detail_response = client.get(f"/api/v1/jobs/{jobs_payload['items'][0]['job_id']}")

        assert detail_response.status_code == 200
        detail_payload = detail_response.json()["data"]
        assert detail_payload["title"] == "Senior ML Engineer"
        assert detail_payload["company_name"] == "Acme AI"
        assert detail_payload["score_breakdown"]["semantic_similarity"] > 0.5

        search_response = client.get("/api/v1/jobs/semantic-search", params={"q": "FastAPI Python remote ML platform"})

        assert search_response.status_code == 200
        search_payload = search_response.json()["data"]
        assert search_payload["total"] == 2
        assert search_payload["items"][0]["title"] == "Senior ML Engineer"
        assert search_payload["items"][1]["title"] == "Lead Data Engineer"


def test_jobs_apply_url_uniqueness_is_enforced(tmp_path: Path) -> None:
    async def run() -> None:
        engine = create_async_engine(f"sqlite+aiosqlite:///{tmp_path / 'jobs-unique.db'}")
        async with engine.begin() as connection:
            await connection.run_sync(Base.metadata.create_all)

        session_factory = async_sessionmaker(engine, expire_on_commit=False)
        async with session_factory() as session:
            first = Job(
                external_id="duplicate-1",
                source="linkedin",
                title="Platform Engineer",
                company_name="Acme AI",
                company_domain="acme.ai",
                location="Remote",
                is_remote=True,
                salary_min=2500000,
                salary_max=3500000,
                description_text="First duplicate job",
                description_embedding=EmbeddingService().embed_text("First duplicate job"),
                apply_url="https://jobs.example.ai/platform-engineer",
            )
            second = Job(
                external_id="duplicate-2",
                source="indeed",
                title="Platform Engineer",
                company_name="Acme AI",
                company_domain="acme.ai",
                location="Remote",
                is_remote=True,
                salary_min=2500000,
                salary_max=3500000,
                description_text="Second duplicate job",
                description_embedding=EmbeddingService().embed_text("Second duplicate job"),
                apply_url="https://jobs.example.ai/platform-engineer",
            )

            session.add(first)
            await session.commit()
            session.add(second)

            with pytest.raises(IntegrityError):
                await session.commit()

        await engine.dispose()

    anyio.run(run)


async def _create_all_tables(engine) -> None:
    async with engine.begin() as connection:
        await connection.run_sync(Base.metadata.create_all)


def _register_and_prepare_resume(client: TestClient) -> None:
    register_response = client.post(
        "/api/v1/auth/register",
        json={
            "email": "jobs-user@example.com",
            "password": "SuperSecret123!",
            "full_name": "Jobs User",
        },
    )

    assert register_response.status_code == 201

    upload_response = client.post(
        "/api/v1/resume/upload",
        files={
            "file": (
                "resume.docx",
                _build_resume_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert upload_response.status_code == 201

    preferences_response = client.put(
        "/api/v1/resume/preferences",
        json={
            "target_roles": ["ML Engineer", "AI Engineer"],
            "preferred_locations": ["Remote", "Bengaluru"],
            "remote_preference": "remote",
            "salary_min": 2500000,
            "salary_max": 4000000,
            "currency": "INR",
            "excluded_companies": ["Example Corp"],
            "seniority_level": "senior",
            "is_active": True,
        },
    )

    assert preferences_response.status_code == 200


def _build_resume_docx() -> bytes:
    document = Document()
    document.add_paragraph("Jobs User")
    document.add_paragraph("jobs-user@example.com")
    document.add_paragraph("Senior ML Engineer")
    document.add_paragraph("Skills: Python, FastAPI, PostgreSQL, Docker, LangGraph")
    document.add_paragraph("Experience")
    document.add_paragraph("Acme AI | Senior ML Engineer | 2021-Present")
    document.add_paragraph("Built an automation pipeline that reduced review time by 42 percent.")
    document.add_paragraph("Education")
    document.add_paragraph("B.Tech Computer Science | ABC University | 2020")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def _seed_jobs(engine) -> None:
    embedding_service = EmbeddingService()
    session_factory = async_sessionmaker(engine, expire_on_commit=False)

    async with session_factory() as session:
        jobs = [
            Job(
                external_id="good-match-1",
                source="linkedin",
                title="Senior ML Engineer",
                company_name="Acme AI",
                company_domain="acme.ai",
                location="Remote",
                is_remote=True,
                salary_min=2600000,
                salary_max=3500000,
                description_text="Python FastAPI PostgreSQL Docker LangGraph machine learning platform remote role.",
                description_embedding=embedding_service.embed_text(
                    "Python FastAPI PostgreSQL Docker LangGraph machine learning platform remote role."
                ),
                apply_url="https://jobs.acme.ai/senior-ml-engineer",
            ),
            Job(
                external_id="close-match-1",
                source="indeed",
                title="Lead Data Engineer",
                company_name="Northstar Labs",
                company_domain="northstar.dev",
                location="Remote",
                is_remote=True,
                salary_min=3000000,
                salary_max=3900000,
                description_text="Python PostgreSQL data pipelines cloud platform engineering role.",
                description_embedding=embedding_service.embed_text(
                    "Python PostgreSQL data pipelines cloud platform engineering role."
                ),
                apply_url="https://jobs.northstar.dev/lead-data-engineer",
            ),
            Job(
                external_id="filtered-location",
                source="remotive",
                title="Senior ML Engineer",
                company_name="City AI",
                company_domain="cityai.example",
                location="Mumbai",
                is_remote=False,
                salary_min=2800000,
                salary_max=3600000,
                description_text="Python FastAPI ML systems onsite role in Mumbai.",
                description_embedding=embedding_service.embed_text(
                    "Python FastAPI ML systems onsite role in Mumbai."
                ),
                apply_url="https://cityai.example/jobs/ml-onsite",
            ),
            Job(
                external_id="filtered-salary",
                source="wellfound",
                title="Senior ML Engineer",
                company_name="Budget AI",
                company_domain="budget.example",
                location="Remote",
                is_remote=True,
                salary_min=1200000,
                salary_max=1800000,
                description_text="Python ML role but below salary preference.",
                description_embedding=embedding_service.embed_text(
                    "Python ML role but below salary preference."
                ),
                apply_url="https://budget.example/jobs/ml",
            ),
            Job(
                external_id="filtered-company",
                source="serpapi",
                title="Senior AI Engineer",
                company_name="Example Corp",
                company_domain="example.com",
                location="Remote",
                is_remote=True,
                salary_min=2800000,
                salary_max=3800000,
                description_text="Strong match but excluded company.",
                description_embedding=embedding_service.embed_text(
                    "Strong match but excluded company."
                ),
                apply_url="https://example.com/jobs/ai-engineer",
            ),
        ]

        session.add_all(jobs)
        await session.commit()
