from __future__ import annotations

from io import BytesIO
from pathlib import Path

import anyio
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import select

from app.core.config import Settings
from app.main import create_app
from app.models.application import Application
from app.models.base import Base
from app.models.job import Job
from app.models.user import User
from app.services.email_monitor_service import EmailMessage, EmailMonitorService


def test_applications_and_notifications_reflect_detected_replies(tmp_path: Path) -> None:
    settings = Settings(
        environment="test",
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'applications.db'}",
        redis_url="redis://localhost:6397/0",
        jwt_secret_key="test-jwt-secret-key-with-32-characters",
        fernet_secret_key="wWKJg6WVKwwhFVWG2yt30YIOCwVDDDeWGPAHDLcGRID=",
        encryption_pepper="pepper-for-tests",
    )

    async def healthy_reporter() -> dict[str, str]:
        return {"status": "ok", "db": "up", "redis": "up"}

    app = create_app(settings=settings, health_reporter=healthy_reporter)
    app.state.redis = _InMemoryRedisClient()

    with TestClient(app) as client:
        anyio.run(_create_all_tables, app.state.database.engine)
        _register_and_prepare_resume(client)
        application_id = _create_applied_application(client)

        async def process_recruiter_reply() -> None:
            async with app.state.database.session() as session:
                user = await session.scalar(select(User).where(User.email == "pipeline-user@example.com"))
                assert user is not None
                application = await session.scalar(select(Application).where(Application.id == application_id))
                assert application is not None
                job = await session.scalar(select(Job).where(Job.id == application.job_id))
                assert job is not None
                service = EmailMonitorService()
                await service.process_messages(
                    session=session,
                    user=user,
                    messages=[
                        EmailMessage(
                            thread_id="gmail-thread-1",
                            sender=f"recruiter@{job.company_domain}",
                            subject=f"Interview availability for {job.company_name}",
                            body="We would love to schedule an interview next week for this role.",
                            snippet="We would love to schedule an interview next week.",
                        )
                    ],
                )

        anyio.run(process_recruiter_reply)

        list_response = client.get("/api/v1/applications")

        assert list_response.status_code == 200
        applications_payload = list_response.json()["data"]["items"]
        assert len(applications_payload) >= 1
        tracked_application = next(item for item in applications_payload if item["id"] == application_id)
        assert tracked_application["status"] == "interview_requested"
        assert tracked_application["latest_email_classification"] == "interview_request"

        detail_response = client.get(f"/api/v1/applications/{application_id}")

        assert detail_response.status_code == 200
        detail_payload = detail_response.json()["data"]
        assert detail_payload["status"] == "interview_requested"
        assert detail_payload["email_monitor"] is not None
        assert detail_payload["email_monitor"]["latest_classification"] == "interview_request"
        assert "interview" in detail_payload["email_monitor"]["snippet"].lower()

        notifications_response = client.get("/api/v1/notifications")

        assert notifications_response.status_code == 200
        assert notifications_response.headers["content-type"].startswith("text/event-stream")
        assert "interview_request" in notifications_response.text

        stats_response = client.get("/api/v1/applications/stats")

        assert stats_response.status_code == 200
        stats_payload = stats_response.json()["data"]
        assert stats_payload["total_applications"] >= 1
        assert stats_payload["total_applied"] >= 1
        assert stats_payload["total_replied"] >= 1
        assert stats_payload["response_rate"] > 0
        assert stats_payload["avg_hours_to_first_reply"] is not None
        assert stats_payload["avg_hours_to_first_reply"] >= 0
        assert any(item["source"] == "indeed" for item in stats_payload["source_breakdown"])
        assert any(item["title"] == tracked_application["title"] for item in stats_payload["top_titles"])


async def _create_all_tables(engine) -> None:
    async with engine.begin() as connection:
        await connection.run_sync(Base.metadata.create_all)


def _create_applied_application(client: TestClient) -> str:
    start_response = client.post(
        "/api/v1/pipeline/start",
        json={
            "target_role": "ML Engineer",
            "location": "Remote",
            "limit_per_source": 10,
            "sources": ["indeed"],
        },
    )

    assert start_response.status_code == 202
    run_id = start_response.json()["data"]["run_id"]

    results_response = client.get(f"/api/v1/pipeline/{run_id}/results")
    assert results_response.status_code == 200
    application_id = results_response.json()["data"]["applications"][0]["id"]

    approve_response = client.post(
        f"/api/v1/pipeline/{run_id}/approve",
        json={"application_ids": [application_id]},
    )

    assert approve_response.status_code == 200
    return application_id


def _register_and_prepare_resume(client: TestClient) -> None:
    register_response = client.post(
        "/api/v1/auth/register",
        json={
            "email": "pipeline-user@example.com",
            "password": "SuperSecret123!",
            "full_name": "Pipeline User",
        },
    )

    assert register_response.status_code == 201

    upload_response = client.post(
        "/api/v1/resume/upload",
        files={
            "file": (
                "resume.docx",
                _build_resume_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert upload_response.status_code == 201

    preferences_response = client.put(
        "/api/v1/resume/preferences",
        json={
            "target_roles": ["ML Engineer", "AI Engineer"],
            "preferred_locations": ["Remote", "Bengaluru"],
            "remote_preference": "remote",
            "salary_min": 2500000,
            "salary_max": 4000000,
            "currency": "INR",
            "excluded_companies": [],
            "seniority_level": "senior",
            "is_active": True,
        },
    )

    assert preferences_response.status_code == 200


def _build_resume_docx() -> bytes:
    document = Document()
    document.add_paragraph("Pipeline User")
    document.add_paragraph("pipeline-user@example.com")
    document.add_paragraph("Senior ML Engineer")
    document.add_paragraph("Skills: Python, FastAPI, PostgreSQL, Docker, LangGraph")
    document.add_paragraph("Experience")
    document.add_paragraph("Acme AI | Senior ML Engineer | 2021-Present")
    document.add_paragraph("Built an automation pipeline that reduced review time by 42 percent.")
    document.add_paragraph("Education")
    document.add_paragraph("B.Tech Computer Science | ABC University | 2020")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class _InMemoryRedisClient:
    def __init__(self) -> None:
        self._data: dict[str, str] = {}
        self._counters: dict[str, int] = {}

    @property
    def client(self) -> "_InMemoryRedisClient":
        return self

    async def incr(self, key: str) -> int:
        self._counters[key] = self._counters.get(key, 0) + 1
        return self._counters[key]

    async def expire(self, key: str, window_seconds: int) -> bool:
        return True

    async def set(self, key: str, value: str, ex: int | None = None) -> bool:
        self._data[key] = value
        return True

    async def get(self, key: str) -> str | None:
        return self._data.get(key)

    async def delete(self, *keys: str) -> int:
        deleted = 0
        for key in keys:
            if key in self._data:
                deleted += 1
                del self._data[key]
        return deleted

    async def scan_iter(self, match: str | None = None):
        prefix = match[:-1] if match and match.endswith("*") else match
        for key in list(self._data.keys()):
            if prefix is None or key.startswith(prefix):
                yield key

    async def ping(self) -> bool:
        return True

    async def close(self) -> None:
        return None
