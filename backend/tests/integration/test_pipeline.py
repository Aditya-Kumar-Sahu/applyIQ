from __future__ import annotations

from io import BytesIO
from pathlib import Path
import re

import anyio
from docx import Document
from fastapi.testclient import TestClient

from app.core.config import Settings
from app.main import create_app
from app.models.agent_run import AgentRun
from app.models.base import Base
from app.models.pipeline_run import PipelineRun


def test_pipeline_start_pause_edit_approve_and_complete(tmp_path: Path) -> None:
    settings = Settings(
        environment="test",
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'pipeline.db'}",
        redis_url="redis://localhost:6395/0",
        jwt_secret_key="test-jwt-secret-key-with-32-characters",
        fernet_secret_key="wWKJg6WVKwwhFVWG2yt30YIOCwVDDDeWGPAHDLcGRID=",
        encryption_pepper="pepper-for-tests",
    )

    async def healthy_reporter() -> dict[str, str]:
        return {"status": "ok", "db": "up", "redis": "up"}

    app = create_app(settings=settings, health_reporter=healthy_reporter)
    app.state.redis = _InMemoryRedisClient()

    with TestClient(app) as client:
        anyio.run(_create_all_tables, app.state.database.engine)
        _register_and_prepare_resume(client)

        start_response = client.post(
            "/api/v1/pipeline/start",
            json={
                "target_role": "ML Engineer",
                "location": "Remote",
                "limit_per_source": 10,
                "sources": ["indeed"],
            },
        )

        assert start_response.status_code == 202
        start_payload = start_response.json()["data"]
        assert start_payload["status"] == "paused_at_gate"
        assert start_payload["current_node"] == "approval_gate_node"
        assert start_payload["pending_approvals_count"] == 5

        run_id = start_payload["run_id"]
        snapshot_key = f"pipeline_run_state:{run_id}"
        checkpoint_key = f"pipeline_run_checkpoint:{run_id}"
        assert anyio.run(_redis_key_exists, app.state.redis, snapshot_key) is True
        assert anyio.run(_redis_key_exists, app.state.redis, checkpoint_key) is True

        status_response = client.get(f"/api/v1/pipeline/{run_id}/status")

        assert status_response.status_code == 200
        assert status_response.headers["content-type"].startswith("text/event-stream")
        assert "approval_gate_node" in status_response.text

        results_response = client.get(f"/api/v1/pipeline/{run_id}/results")

        assert results_response.status_code == 200
        results_payload = results_response.json()["data"]
        assert results_payload["status"] == "paused_at_gate"
        assert len(results_payload["applications"]) == 5
        assert len({application["cover_letter_text"] for application in results_payload["applications"]}) == 5

        for application in results_payload["applications"]:
            assert application["company_name"] in application["cover_letter_text"]
            assert application["word_count"] <= 250
            assert application["tone"] in {"formal", "conversational"}
            assert "I am writing to express my interest" not in application["cover_letter_text"]
            assert "I believe I would be a great fit" not in application["cover_letter_text"]
            assert "Please find attached" not in application["cover_letter_text"]
            assert re.search(r"\d", application["cover_letter_text"])

        application_id = results_payload["applications"][0]["id"]
        original_cover_letter = results_payload["applications"][0]["cover_letter_text"]
        rejected_application_id = results_payload["applications"][1]["id"]

        regenerate_response = client.post(
            f"/api/v1/pipeline/{run_id}/application/{application_id}/cover-letter/regenerate"
        )

        assert regenerate_response.status_code == 200
        regenerated_payload = regenerate_response.json()["data"]
        assert regenerated_payload["cover_letter_version"] == 2
        assert regenerated_payload["cover_letter_text"] != original_cover_letter
        assert regenerated_payload["tone"] in {"formal", "conversational"}

        edit_response = client.put(
            f"/api/v1/pipeline/{run_id}/application/{application_id}/cover-letter",
            json={"cover_letter_text": "A sharper, edited cover letter for the approval gate."},
        )

        assert edit_response.status_code == 200
        assert edit_response.json()["data"]["cover_letter_version"] == 3
        assert edit_response.json()["data"]["tone"] == "edited"

        ab_test_response = client.post(
            f"/api/v1/pipeline/{run_id}/application/{application_id}/cover-letter/ab-test"
        )

        assert ab_test_response.status_code == 200
        ab_test_payload = ab_test_response.json()["data"]
        assert ab_test_payload["application_id"] == application_id
        assert ab_test_payload["cover_letter_version"] == 3
        assert len(ab_test_payload["variants"]) == 2
        assert {variant["variant_id"] for variant in ab_test_payload["variants"]} == {"A", "B"}
        assert ab_test_payload["variants"][0]["cover_letter_text"] != ab_test_payload["variants"][1]["cover_letter_text"]
        assert all(variant["word_count"] <= 250 for variant in ab_test_payload["variants"])

        select_variant_response = client.post(
            f"/api/v1/pipeline/{run_id}/application/{application_id}/cover-letter/select-variant",
            json={"variant_id": "B"},
        )

        assert select_variant_response.status_code == 200
        select_variant_payload = select_variant_response.json()["data"]
        assert select_variant_payload["application_id"] == application_id
        assert select_variant_payload["selected_variant_id"] == "B"
        assert select_variant_payload["cover_letter_version"] == 4
        assert select_variant_payload["tone"] in {"formal", "conversational"}

        reject_response = client.post(
            f"/api/v1/pipeline/{run_id}/reject",
            json={"application_ids": [rejected_application_id]},
        )

        assert reject_response.status_code == 200
        assert reject_response.json()["data"]["rejected_count"] == 1

        approve_response = client.post(
            f"/api/v1/pipeline/{run_id}/approve",
            json={"application_ids": [application_id]},
        )

        assert approve_response.status_code == 200
        approve_payload = approve_response.json()["data"]
        assert approve_payload["status"] == "complete"
        assert approve_payload["applications_submitted"] == 1
        assert anyio.run(_redis_key_exists, app.state.redis, snapshot_key) is False
        assert anyio.run(_redis_key_exists, app.state.redis, checkpoint_key) is False

        completed_results = client.get(f"/api/v1/pipeline/{run_id}/results")

        assert completed_results.status_code == 200
        completed_payload = completed_results.json()["data"]
        assert completed_payload["status"] == "complete"
        applied_application = next(app for app in completed_payload["applications"] if app["id"] == application_id)
        assert applied_application["status"] == "applied"
        assert applied_application["ats_provider"] == "indeed_apply"
        assert applied_application["confirmation_url"]
        assert len(applied_application["screenshot_urls"]) == 2
        assert applied_application["selected_variant_id"] == "B"
        untouched_applications = [
            app for app in completed_payload["applications"] if app["status"] == "pending_approval"
        ]
        assert len(untouched_applications) == 3
        assert any(app["status"] == "rejected" for app in completed_payload["applications"])

        agent_run_count = anyio.run(_count_agent_runs, app.state.database.engine)
        assert agent_run_count >= 5


def test_pipeline_start_reuses_existing_active_run(tmp_path: Path) -> None:
    settings = Settings(
        environment="test",
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'pipeline-reuse.db'}",
        redis_url="redis://localhost:6395/0",
        jwt_secret_key="test-jwt-secret-key-with-32-characters",
        fernet_secret_key="wWKJg6WVKwwhFVWG2yt30YIOCwVDDDeWGPAHDLcGRID=",
        encryption_pepper="pepper-for-tests",
    )

    async def healthy_reporter() -> dict[str, str]:
        return {"status": "ok", "db": "up", "redis": "up"}

    app = create_app(settings=settings, health_reporter=healthy_reporter)
    app.state.redis = _InMemoryRedisClient()

    with TestClient(app) as client:
        anyio.run(_create_all_tables, app.state.database.engine)
        _register_and_prepare_resume(client)

        first_response = client.post(
            "/api/v1/pipeline/start",
            json={
                "target_role": "ML Engineer",
                "location": "Remote",
                "limit_per_source": 10,
                "sources": ["indeed"],
            },
        )

        assert first_response.status_code == 202
        first_payload = first_response.json()["data"]
        assert first_payload["status"] == "paused_at_gate"

        second_response = client.post(
            "/api/v1/pipeline/start",
            json={
                "target_role": "ML Engineer",
                "location": "Remote",
                "limit_per_source": 10,
                "sources": ["linkedin", "indeed", "remotive"],
            },
        )

        assert second_response.status_code == 202
        second_payload = second_response.json()["data"]
        assert second_payload["run_id"] == first_payload["run_id"]
        assert second_payload["status"] == "paused_at_gate"
        assert second_payload["current_node"] == "approval_gate_node"
        assert second_payload["pending_approvals_count"] == first_payload["pending_approvals_count"]

        run_count = anyio.run(_count_pipeline_runs, app.state.database.engine)
        assert run_count == 1


def test_delete_account_purges_pipeline_redis_state(tmp_path: Path) -> None:
    settings = Settings(
        environment="test",
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'pipeline-delete.db'}",
        redis_url="redis://localhost:6395/0",
        jwt_secret_key="test-jwt-secret-key-with-32-characters",
        fernet_secret_key="wWKJg6WVKwwhFVWG2yt30YIOCwVDDDeWGPAHDLcGRID=",
        encryption_pepper="pepper-for-tests",
    )

    async def healthy_reporter() -> dict[str, str]:
        return {"status": "ok", "db": "up", "redis": "up"}

    app = create_app(settings=settings, health_reporter=healthy_reporter)
    app.state.redis = _InMemoryRedisClient()

    with TestClient(app) as client:
        anyio.run(_create_all_tables, app.state.database.engine)
        _register_and_prepare_resume(client)

        start_response = client.post(
            "/api/v1/pipeline/start",
            json={
                "target_role": "ML Engineer",
                "location": "Remote",
                "limit_per_source": 10,
                "sources": ["indeed"],
            },
        )

        assert start_response.status_code == 202
        run_id = start_response.json()["data"]["run_id"]

        snapshot_key = f"pipeline_run_state:{run_id}"
        checkpoint_key = f"pipeline_run_checkpoint:{run_id}"
        assert anyio.run(_redis_key_exists, app.state.redis, snapshot_key) is True
        assert anyio.run(_redis_key_exists, app.state.redis, checkpoint_key) is True

        delete_response = client.request(
            "DELETE",
            "/api/v1/auth/account",
            json={"password_confirmation": "SuperSecret123!"},
        )

        assert delete_response.status_code == 200
        assert anyio.run(_redis_key_exists, app.state.redis, snapshot_key) is False
        assert anyio.run(_redis_key_exists, app.state.redis, checkpoint_key) is False


async def _create_all_tables(engine) -> None:
    async with engine.begin() as connection:
        await connection.run_sync(Base.metadata.create_all)


async def _count_agent_runs(engine) -> int:
    from sqlalchemy.ext.asyncio import async_sessionmaker
    from sqlalchemy import select, func

    session_factory = async_sessionmaker(engine, expire_on_commit=False)
    async with session_factory() as session:
        result = await session.scalar(select(func.count()).select_from(AgentRun))
        return int(result or 0)


async def _count_pipeline_runs(engine) -> int:
    from sqlalchemy.ext.asyncio import async_sessionmaker
    from sqlalchemy import select, func

    session_factory = async_sessionmaker(engine, expire_on_commit=False)
    async with session_factory() as session:
        result = await session.scalar(select(func.count()).select_from(PipelineRun))
        return int(result or 0)


async def _redis_key_exists(redis_client, key: str) -> bool:
    return (await redis_client.get(key)) is not None


class _InMemoryRedisClient:
    def __init__(self) -> None:
        self._data: dict[str, str] = {}
        self._counters: dict[str, int] = {}

    @property
    def client(self) -> "_InMemoryRedisClient":
        return self

    async def incr(self, key: str) -> int:
        self._counters[key] = self._counters.get(key, 0) + 1
        return self._counters[key]

    async def expire(self, key: str, window_seconds: int) -> bool:
        return True

    async def set(self, key: str, value: str, ex: int | None = None) -> bool:
        self._data[key] = value
        return True

    async def get(self, key: str) -> str | None:
        return self._data.get(key)

    async def delete(self, *keys: str) -> int:
        deleted = 0
        for key in keys:
            if key in self._data:
                deleted += 1
                del self._data[key]
        return deleted

    async def scan_iter(self, match: str | None = None):
        prefix = match[:-1] if match and match.endswith("*") else match
        for key in list(self._data.keys()):
            if prefix is None or key.startswith(prefix):
                yield key

    async def ping(self) -> bool:
        return True

    async def close(self) -> None:
        return None


def _register_and_prepare_resume(client: TestClient) -> None:
    register_response = client.post(
        "/api/v1/auth/register",
        json={
            "email": "pipeline-user@example.com",
            "password": "SuperSecret123!",
            "full_name": "Pipeline User",
        },
    )

    assert register_response.status_code == 201

    upload_response = client.post(
        "/api/v1/resume/upload",
        files={
            "file": (
                "resume.docx",
                _build_resume_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert upload_response.status_code == 201

    preferences_response = client.put(
        "/api/v1/resume/preferences",
        json={
            "target_roles": ["ML Engineer", "AI Engineer"],
            "preferred_locations": ["Remote", "Bengaluru"],
            "remote_preference": "remote",
            "salary_min": 2500000,
            "salary_max": 4000000,
            "currency": "INR",
            "excluded_companies": [],
            "seniority_level": "senior",
            "is_active": True,
        },
    )

    assert preferences_response.status_code == 200


def _build_resume_docx() -> bytes:
    document = Document()
    document.add_paragraph("Pipeline User")
    document.add_paragraph("pipeline-user@example.com")
    document.add_paragraph("Senior ML Engineer")
    document.add_paragraph("Skills: Python, FastAPI, PostgreSQL, Docker, LangGraph")
    document.add_paragraph("Experience")
    document.add_paragraph("Acme AI | Senior ML Engineer | 2021-Present")
    document.add_paragraph("Built an automation pipeline that reduced review time by 42 percent.")
    document.add_paragraph("Education")
    document.add_paragraph("B.Tech Computer Science | ABC University | 2020")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
