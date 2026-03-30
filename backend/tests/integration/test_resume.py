from __future__ import annotations

from io import BytesIO
from pathlib import Path

import anyio
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import select
from sqlalchemy.ext.asyncio import async_sessionmaker

from app.core.config import Settings
from app.main import create_app
from app.models.base import Base
from app.models.resume_profile import ResumeProfile


def test_resume_upload_profile_preferences_and_completeness(tmp_path: Path) -> None:
    settings = Settings(
        environment="test",
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'resume.db'}",
        redis_url="redis://localhost:6391/0",
        jwt_secret_key="test-jwt-secret-key-with-32-characters",
        fernet_secret_key="wWKJg6WVKwwhFVWG2yt30YIOCwVDDDeWGPAHDLcGRID=",
        encryption_pepper="pepper-for-tests",
    )

    async def healthy_reporter() -> dict[str, str]:
        return {"status": "ok", "db": "up", "redis": "up"}

    app = create_app(settings=settings, health_reporter=healthy_reporter)

    with TestClient(app) as client:
        anyio.run(_create_all_tables, app.state.database.engine)

        register_response = client.post(
            "/api/v1/auth/register",
            json={
                "email": "resume-user@example.com",
                "password": "SuperSecret123!",
                "full_name": "Resume User",
            },
        )

        assert register_response.status_code == 201

        upload_response = client.post(
            "/api/v1/resume/upload",
            files={
                "file": (
                    "resume.docx",
                    _build_resume_docx(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        assert upload_response.status_code == 201
        assert upload_response.json()["data"]["profile"]["current_title"] == "Senior ML Engineer"
        assert "Python" in upload_response.json()["data"]["profile"]["skills"]["technical"]

        resume_response = client.get("/api/v1/resume")

        assert resume_response.status_code == 200
        assert resume_response.json()["data"]["profile"]["email"] == "resume-user@example.com"
        assert resume_response.json()["data"]["preferences"] is None

        preferences_response = client.put(
            "/api/v1/resume/preferences",
            json={
                "target_roles": ["ML Engineer", "AI Engineer"],
                "preferred_locations": ["Remote", "Bengaluru"],
                "remote_preference": "remote",
                "salary_min": 2500000,
                "salary_max": 4000000,
                "currency": "INR",
                "excluded_companies": ["Example Corp"],
                "seniority_level": "senior",
                "is_active": True,
            },
        )

        assert preferences_response.status_code == 200
        assert preferences_response.json()["data"]["preferences"]["target_roles"] == ["ML Engineer", "AI Engineer"]

        completeness_response = client.get("/api/v1/resume/profile-completeness")

        assert completeness_response.status_code == 200
        assert completeness_response.json()["data"]["score"] >= 80
        assert completeness_response.json()["data"]["missing_fields"] == []

        stored_resume = anyio.run(_get_stored_resume, app.state.database.engine)

        assert stored_resume is not None
        assert "Senior ML Engineer" not in stored_resume.raw_text
        assert stored_resume.file_hash
        assert len(stored_resume.resume_embedding) > 0


def test_resume_upload_rejects_spoofed_extension(tmp_path: Path) -> None:
    settings = Settings(
        environment="test",
        database_url=f"sqlite+aiosqlite:///{tmp_path / 'resume_spoof.db'}",
        redis_url="redis://localhost:6391/0",
        jwt_secret_key="test-jwt-secret-key-with-32-characters",
        fernet_secret_key="wWKJg6WVKwwhFVWG2yt30YIOCwVDDDeWGPAHDLcGRID=",
        encryption_pepper="pepper-for-tests",
    )

    async def healthy_reporter() -> dict[str, str]:
        return {"status": "ok", "db": "up", "redis": "up"}

    app = create_app(settings=settings, health_reporter=healthy_reporter)

    with TestClient(app) as client:
        anyio.run(_create_all_tables, app.state.database.engine)

        register_response = client.post(
            "/api/v1/auth/register",
            json={
                "email": "spoof-user@example.com",
                "password": "SuperSecret123!",
                "full_name": "Spoof User",
            },
        )

        assert register_response.status_code == 201

        spoofed_upload_response = client.post(
            "/api/v1/resume/upload",
            files={
                "file": (
                    "resume.pdf",
                    _build_resume_docx(),
                    "application/pdf",
                )
            },
        )

        assert spoofed_upload_response.status_code == 400
        assert spoofed_upload_response.json()["error"]["message"] == "Resume file content does not match the declared file extension"


def _build_resume_docx() -> bytes:
    document = Document()
    document.add_paragraph("Resume User")
    document.add_paragraph("resume-user@example.com")
    document.add_paragraph("Senior ML Engineer")
    document.add_paragraph("Skills: Python, FastAPI, PostgreSQL, Docker, LangGraph")
    document.add_paragraph("Experience")
    document.add_paragraph("Acme AI | Senior ML Engineer | 2021-Present")
    document.add_paragraph("Built an automation pipeline that reduced review time by 42 percent.")
    document.add_paragraph("Education")
    document.add_paragraph("B.Tech Computer Science | ABC University | 2020")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def _create_all_tables(engine) -> None:
    async with engine.begin() as connection:
        await connection.run_sync(Base.metadata.create_all)


async def _get_stored_resume(engine) -> ResumeProfile | None:
    session_factory = async_sessionmaker(engine, expire_on_commit=False)
    async with session_factory() as session:
        return await session.scalar(select(ResumeProfile))
