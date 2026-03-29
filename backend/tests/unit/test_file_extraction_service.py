from __future__ import annotations

from io import BytesIO

import fitz
from docx import Document

from app.services.file_extraction_service import FileExtractionService


def test_extract_docx_text_includes_paragraphs_and_tables() -> None:
    service = FileExtractionService()
    document = Document()
    document.add_paragraph("Aditya Sahu")
    document.add_paragraph("Skills: Python, FastAPI")
    table = document.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "Acme AI"
    table.cell(0, 1).text = "Senior Backend Engineer"
    table.cell(0, 2).text = "2021 - Present"

    buffer = BytesIO()
    document.save(buffer)

    extracted = service.extract_text("resume.docx", buffer.getvalue())

    assert "Aditya Sahu" in extracted
    assert "Skills: Python, FastAPI" in extracted
    assert "Acme AI | Senior Backend Engineer | 2021 - Present" in extracted


def test_extract_pdf_text_reads_inserted_page_text() -> None:
    service = FileExtractionService()
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Senior Backend Engineer")
    page.insert_text((72, 90), "Python, FastAPI, PostgreSQL")
    content = pdf.tobytes()

    extracted = service.extract_text("resume.pdf", content)

    assert "Senior Backend Engineer" in extracted
    assert "Python, FastAPI, PostgreSQL" in extracted
